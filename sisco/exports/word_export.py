"""
Exportación a Word con python-docx
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from sisco.database.models import Sesion, TipoPuntoEnum


MAGENTA = RGBColor(0xE9, 0x1E, 0x8C)


def _heading(doc: Document, texto: str, nivel: int = 1):
    p = doc.add_heading(texto, level=nivel)
    for run in p.runs:
        run.font.color.rgb = MAGENTA


def exportar_sesion_word(sesion: Sesion) -> bytes:
    doc = Document()

    # Título
    titulo = doc.add_heading("", 0)
    run = titulo.add_run("SISCO — Reporte de Sesión")
    run.font.color.rgb = MAGENTA
    run.font.size = Pt(18)

    # Datos generales
    _heading(doc, "Datos generales", 1)
    tabla_gen = doc.add_table(rows=5, cols=2)
    tabla_gen.style = "Table Grid"
    datos = [
        ("Naturaleza", sesion.naturaleza.value),
        ("Tipo de sesión", sesion.tipo_sesion),
        ("Fecha", str(sesion.fecha)),
        ("No. de sesión", str(sesion.numero_sesion)),
        ("Modalidad", sesion.modalidad.value),
    ]
    for i, (campo, valor) in enumerate(datos):
        tabla_gen.cell(i, 0).text = campo
        tabla_gen.cell(i, 1).text = valor

    # Horas
    _heading(doc, "Horas", 2)
    doc.add_paragraph(f"Inicio transmisión: {sesion.hora_transmision or '—'}")
    doc.add_paragraph(f"Instalación: {sesion.hora_instalacion or '—'}")
    doc.add_paragraph(f"Clausura: {sesion.hora_clausura or '—'}")

    # Asistencia
    _heading(doc, "Asistencia", 1)
    if sesion.asistencias:
        tabla_a = doc.add_table(rows=1, cols=4)
        tabla_a.style = "Table Grid"
        for i, h in enumerate(["Cargo", "Nombre", "Asistencia", "Modalidad"]):
            tabla_a.cell(0, i).text = h
        for a in sesion.asistencias:
            row = tabla_a.add_row()
            row.cells[0].text = a.integrante.cargo.value
            row.cells[1].text = a.integrante.nombre
            row.cells[2].text = "Presente" if a.presente else f"Falta — {a.justificacion or ''}"
            row.cells[3].text = a.modalidad.value if a.modalidad else "—"
    else:
        doc.add_paragraph("Sin registro de asistencia.")

    # Puntos del orden del día
    _heading(doc, "Orden del día", 1)
    puntos = [p for p in sesion.puntos if p.incluido]
    for punto in puntos:
        _heading(doc, f"Punto {punto.orden}: {punto.tipo.value}", 2)
        if punto.descripcion:
            doc.add_paragraph(punto.descripcion)

        # Participaciones
        if punto.participaciones:
            doc.add_paragraph("Participaciones:", style="Intense Quote")
            for part in punto.participaciones:
                doc.add_paragraph(
                    f"[Ronda {part.ronda or '—'}] {part.integrante.nombre} ({part.tipo.value})"
                    + (f": {part.texto}" if part.texto else ""),
                    style="List Bullet",
                )

        # Modificaciones
        if punto.modificaciones:
            doc.add_paragraph("Modificaciones:", style="Intense Quote")
            for mod in punto.modificaciones:
                doc.add_paragraph(mod.texto, style="List Bullet")

        # Votación
        if punto.votacion:
            v = punto.votacion
            doc.add_paragraph(
                f"Votación: {v.resultado.value}"
                + (f" — {v.hora_aprobacion}" if v.hora_aprobacion else "")
                + (f"\nObservaciones: {v.observaciones}" if v.observaciones else ""),
                style="Intense Quote",
            )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
